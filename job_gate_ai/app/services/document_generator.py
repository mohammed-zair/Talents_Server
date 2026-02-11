
import os 
from typing import Any ,Dict ,List 
import tempfile 
from datetime import datetime 

try :
    from docx import Document 
    from docx .shared import Inches ,Pt ,RGBColor 
    from docx .enum .text import WD_ALIGN_PARAGRAPH 
    DOCX_AVAILABLE =True 
except ImportError :
    DOCX_AVAILABLE =False 

try :
    from weasyprint import HTML 
    PDF_AVAILABLE =True 
except ImportError :
    PDF_AVAILABLE =False 

class DocumentGenerator :
    def __init__ (self ,output_dir :str ="generated_cvs"):
        self .output_dir =output_dir 
        os .makedirs (output_dir ,exist_ok =True )

    async def generate_pdf (self ,cv_data :Dict [str ,Any ],language :str ="arabic")->str :
        if not PDF_AVAILABLE :
            raise ImportError ("WeasyPrint not installed for PDF generation")


        html_content =self ._create_html_template (cv_data ,language )


        filename =f"cv_{cv_data .get ('personal_info',{}).get ('name') or cv_data .get ('personal_info',{}).get ('full_name','unknown')}_{datetime .now ().strftime ('%Y%m%d_%H%M%S')}.pdf"
        filepath =os .path .join (self .output_dir ,filename )


        HTML (string =html_content ).write_pdf (filepath )

        return filepath 

    async def generate_docx (self ,cv_data :Dict [str ,Any ],language :str ="arabic")->str :
        if not DOCX_AVAILABLE :
            raise ImportError ("python-docx not installed for DOCX generation")


        doc =Document ()


        labels =self ._get_labels (language )
        title =doc .add_heading (labels ["title"],0 )
        title .alignment =WD_ALIGN_PARAGRAPH .CENTER 


        self ._add_personal_info (doc ,cv_data .get ('personal_info',{}))


        self ._add_section (doc ,labels ["summary"],cv_data .get ('summary_professional',''))


        self ._add_experience (doc ,cv_data .get ('experience',[]),labels )


        self ._add_education (doc ,cv_data .get ('education',[]),labels )


        self ._add_skills (doc ,cv_data .get ('skills',[]),labels )


        filename =f"cv_{cv_data .get ('personal_info',{}).get ('name') or cv_data .get ('personal_info',{}).get ('full_name','unknown')}_{datetime .now ().strftime ('%Y%m%d_%H%M%S')}.docx"
        filepath =os .path .join (self .output_dir ,filename )


        doc .save (filepath )

        return filepath 

    async def generate_both (self ,cv_data :Dict [str ,Any ],language :str ="arabic")->Dict [str ,str ]:
        pdf_path =await self .generate_pdf (cv_data ,language )
        docx_path =await self .generate_docx (cv_data ,language )

        return {
        "pdf":pdf_path ,
        "docx":docx_path 
        }

    def generate_html (self ,cv_data :Dict [str ,Any ],language :str ="arabic")->str :
        return self ._create_html_template (cv_data ,language )

    def _create_html_template (self ,cv_data :Dict [str ,Any ],language :str ="arabic")->str :

        personal =cv_data .get ('personal_info',{})
        labels =self ._get_labels (language )
        is_ar =language .lower ()in ["arabic","ar"]
        doc_dir ="rtl"if is_ar else "ltr"
        lang_code ="ar"if is_ar else "en"

        html =f"""
        <!DOCTYPE html>
        <html dir="{doc_dir}" lang="{lang_code}">
        <head>
            <meta charset="UTF-8">
            <title>{labels ["title"]} - {personal .get ('name') or personal .get ('full_name','')}</title>
            <style>
                body {{ font-family: 'Arial', sans-serif; margin: 40px; direction: {doc_dir}; }}
                .header {{ text-align: center; border-bottom: 2px solid #333; padding-bottom: 20px; }}
                .name {{ font-size: 24px; font-weight: bold; }}
                .contact-info {{ margin-top: 10px; }}
                .section {{ margin-top: 30px; }}
                .section-title {{ font-size: 18px; font-weight: bold; border-bottom: 1px solid #ccc; padding-bottom: 5px; }}
                .item {{ margin-top: 15px; }}
                .item-title {{ font-weight: bold; }}
                .skills {{ margin-top: 10px; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="name">{personal .get ('name') or personal .get ('full_name','')}</div>
                <div class="contact-info">
                    {personal .get ('email','')} | {personal .get ('phone','')} | {personal .get ('location','')}
                </div>
            </div>
            
            <div class="section">
                <div class="section-title">{labels ["summary"]}</div>
                <div>{cv_data .get ('summary_professional','') or cv_data .get ('summary','')}</div>
            </div>
            
            <div class="section">
                <div class="section-title">{labels ["experience"]}</div>
                {"".join ([f'<div class="item"><div class="item-title">{exp .get ("position","")} - {exp .get ("company","")}</div><div>{exp .get ("description","")}</div></div>'for exp in cv_data .get ('experience',[])])}
            </div>
            
            <div class="section">
                <div class="section-title">{labels ["education"]}</div>
                {"".join ([f'<div class="item"><div class="item-title">{edu .get ("degree","")} - {edu .get ("institution","")}</div><div>{edu .get ("duration","")}</div></div>'for edu in cv_data .get ('education',[])])}
            </div>
            
            <div class="section">
                <div class="section-title">{labels ["skills"]}</div>
                <div class="skills">{', '.join (cv_data .get ('skills',[]))}</div>
            </div>
        </body>
        </html>
        """

        return html 

    def _add_personal_info (self ,doc ,personal_info :Dict ):
        if personal_info .get ('name') or personal_info .get ('full_name'):
            name =doc .add_paragraph (personal_info .get ('name') or personal_info .get ('full_name'))
            name .runs [0 ].bold =True 
            name .runs [0 ].font .size =Pt (16 )
            name .alignment =WD_ALIGN_PARAGRAPH .CENTER 

        contact =doc .add_paragraph ()
        if personal_info .get ('email'):
            contact .add_run (f"?? {personal_info ['email']} | ")
        if personal_info .get ('phone'):
            contact .add_run (f"?? {personal_info ['phone']} | ")
        if personal_info .get ('location'):
            contact .add_run (f"?? {personal_info ['location']}")
        contact .alignment =WD_ALIGN_PARAGRAPH .CENTER 

        doc .add_paragraph ()

    def _add_section (self ,doc ,title :str ,content :str ):
        if content :
            doc .add_heading (title ,level =1 )
            doc .add_paragraph (content )
            doc .add_paragraph ()

    def _add_experience (self ,doc ,experience :List [Dict ],labels :Dict [str ,str ]):
        if experience :
            doc .add_heading (labels ["experience"],level =1 )
            for exp in experience :
                p =doc .add_paragraph ()
                p .add_run (f"{exp .get ('position','')} - ").bold =True 
                p .add_run (f"{exp .get ('company','')} | {exp .get ('duration','')}")

                if exp .get ('description'):
                    doc .add_paragraph (exp ['description'])

                doc .add_paragraph ()

    def _add_education (self ,doc ,education :List [Dict ],labels :Dict [str ,str ]):
        if education :
            doc .add_heading (labels ["education"],level =1 )
            for edu in education :
                p =doc .add_paragraph ()
                p .add_run (f"{edu .get ('degree','')} - ").bold =True 
                p .add_run (f"{edu .get ('institution','')} | {edu .get ('duration','')}")

                if edu .get ('description'):
                    doc .add_paragraph (edu ['description'])

                doc .add_paragraph ()

    def _add_skills (self ,doc ,skills :List [str ],labels :Dict [str ,str ]):
        if skills :
            doc .add_heading (labels ["skills"],level =1 )
            skills_text =' ? '.join (skills )
            doc .add_paragraph (skills_text )

    def _get_labels (self ,language :str )->Dict [str ,str ]:
        if language .lower ()in ["english","en"]:
            return {
            "title":"Curriculum Vitae",
            "summary":"Professional Summary",
            "experience":"Work Experience",
            "education":"Education",
            "skills":"Skills"
            }
        return {
        "title":"?????? ???????",
        "summary":"???? ???????",
        "experience":"?????? ???????",
        "education":"???????",
        "skills":"????????"
        }
